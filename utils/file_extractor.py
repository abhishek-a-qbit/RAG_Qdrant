import os
import asyncio
import tempfile
from typing import Optional
from fastapi import HTTPException
from services.logger import setup_logger

# Initialize logger
logger = setup_logger("file_extractor")

async def extract_text_from_file(file_content: bytes, file_type: str) -> str:
    """
    Extract text from different file types based on the file type.
    
    Args:
        file_content: Raw bytes of the file
        file_type: Type of file (txt, pdf, docx)
    
    Returns:
        Extracted text as string
    
    Raises:
        HTTPException: If file type is unsupported
    """
    try:
        if file_type == "txt":
            return await extract_text_from_txt(file_content)
        elif file_type == "pdf":
            return await extract_text_from_pdf(file_content)
        elif file_type == "docx":
            return await extract_text_from_docx(file_content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
    except Exception as e:
        logger.error(f"Error extracting text from {file_type} file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")

async def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from TXT file content.
    
    Args:
        file_content: Raw bytes of the TXT file
    
    Returns:
        Extracted text as string
    """
    try:
        # Try UTF-8 first, then fallback to other encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                logger.info(f"Successfully decoded TXT file using {encoding}")
                return text
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use utf-8 with error handling
        text = file_content.decode('utf-8', errors='replace')
        logger.warning("TXT file decoded with UTF-8 error replacement")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from TXT file: {str(e)}")
        raise ValueError(f"Failed to extract text from TXT file: {str(e)}")

async def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file asynchronously.
    
    Args:
        file_content: Raw bytes of the PDF file
    
    Returns:
        Extracted text as string
    """
    try:
        # Check if PyPDF2 is available
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF extraction. Install with: pip install PyPDF2")
        
        # Use asyncio.to_thread to run the blocking PDF extraction in a separate thread
        return await asyncio.to_thread(extract_text_from_pdf_sync, file_content)
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF file: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF file: {str(e)}")

def extract_text_from_pdf_sync(file_content: bytes) -> str:
    """
    Extract text from a PDF file (blocking version).
    
    Args:
        file_content: Raw bytes of the PDF file
    
    Returns:
        Extracted text as string
    """
    try:
        import PyPDF2
        from io import BytesIO
        
        # Create PDF reader from bytes
        pdf_stream = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        
        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            try:
                pdf_reader.decrypt('')
            except:
                logger.warning("PDF is encrypted and requires password")
                return ""
        
        content = ""
        num_pages = len(pdf_reader.pages)
        
        logger.info(f"Processing PDF with {num_pages} pages")
        
        for i in range(num_pages):
            try:
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                if page_text.strip():
                    content += page_text + "\n"
                logger.debug(f"Extracted text from page {i + 1}")
            except Exception as e:
                logger.warning(f"Failed to extract text from page {i + 1}: {str(e)}")
                continue
        
        if not content.strip():
            logger.warning("No text extracted from PDF - might be image-based PDF")
            return ""
        
        logger.info(f"Successfully extracted {len(content)} characters from PDF")
        return content.strip()
        
    except Exception as e:
        logger.error(f"Error in synchronous PDF extraction: {str(e)}")
        raise

async def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file content.
    
    Args:
        file_content: Raw bytes of the DOCX file
    
    Returns:
        Extracted text as string
    """
    try:
        # Check if python-docx is available
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text from DOCX
            doc = docx.Document(temp_file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"HEADER: {paragraph.text}")
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(f"FOOTER: {paragraph.text}")
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                logger.warning("No text extracted from DOCX file")
                return ""
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text.strip()
            
        finally:
            # Clean up temporary file
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
                
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX file: {str(e)}")

def get_supported_file_types() -> list:
    """
    Get list of supported file types.
    
    Returns:
        List of supported file extensions
    """
    return ["txt", "pdf", "docx"]

def is_file_type_supported(file_type: str) -> bool:
    """
    Check if file type is supported.
    
    Args:
        file_type: File extension to check
    
    Returns:
        True if supported, False otherwise
    """
    return file_type.lower() in get_supported_file_types()

def get_file_type_from_filename(filename: str) -> str:
    """
    Extract file type from filename.
    
    Args:
        filename: Name of the file
    
    Returns:
        File extension (lowercase, without dot)
    """
    if '.' not in filename:
        return ""
    
    return filename.split('.')[-1].lower()

async def extract_text_with_metadata(file_content: bytes, filename: str) -> dict:
    """
    Extract text along with metadata from file.
    
    Args:
        file_content: Raw bytes of the file
        filename: Name of the file
    
    Returns:
        Dictionary containing text and metadata
    """
    file_type = get_file_type_from_filename(filename)
    
    if not is_file_type_supported(file_type):
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type: {file_type}. Supported types: {get_supported_file_types()}"
        )
    
    text = await extract_text_from_file(file_content, file_type)
    
    return {
        "text": text,
        "filename": filename,
        "file_type": file_type,
        "file_size": len(file_content),
        "character_count": len(text),
        "word_count": len(text.split()) if text else 0,
        "line_count": len(text.split('\n')) if text else 0
    }
